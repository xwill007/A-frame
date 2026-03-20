#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para convertir MANUAL_USUARIO_BASE.md a documento Word (.docx)
Utiliza python-docx para crear un documento profesional
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def add_horizontal_line(doc):
    """Agrega una línea horizontal (separador)"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_markdown_to_docx(md_file, docx_file):
    """
    Convierte un archivo Markdown a un documento Word (.docx)
    """
    
    # Leer archivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Procesar contenido línea por línea
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Saltar líneas vacías
        if not line.strip():
            i += 1
            continue
        
        # Encabezados (# ## ### etc)
        if line.startswith('#'):
            # Contar número de #
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            # Agregar como encabezado
            if level == 1:
                heading = doc.add_heading(text, level=0)  # Título principal
                heading.style = 'Heading 1'
            elif level == 2:
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
        
        # Línea de separación (---)
        elif line.strip() == '---':
            add_horizontal_line(doc)
        
        # Tablas (| ... | ... |)
        elif line.strip().startswith('|'):
            # Buscar filas de tabla
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j])
                j += 1
            
            if table_lines:
                # Procesar tabla
                rows = [line.split('|')[1:-1] for line in table_lines]
                rows = [[cell.strip() for cell in row] for row in rows]
                
                # Si la segunda fila es separador, saltarla
                if all(all(c in ['-', ':'] for c in cell.strip()) for cell in rows[1 if len(rows) > 1 else 0]):
                    if len(rows) > 1:
                        rows = rows[0:1] + rows[2:]
                
                # Crear tabla en Word
                if rows:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Rellenar tabla
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_data in enumerate(row_data):
                            if c_idx < num_cols:
                                cell = table.rows[r_idx].cells[c_idx]
                                cell.text = cell_data
                                
                                # Hacer negrita encabezados
                                if r_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                
                i = j - 1
        
        # Código (```...```)
        elif line.strip().startswith('```'):
            # Buscar bloque de código
            code_lines = []
            j = i + 1
            while j < len(lines) and not lines[j].strip().startswith('```'):
                code_lines.append(lines[j])
                j += 1
            
            # Agregar código
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'List Bullet'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            i = j
        
        # Listas con viñetas (- o *)
        elif line.strip().startswith(('-', '*')) and not line.strip().startswith('---'):
            p = doc.add_paragraph(line.strip().lstrip('-').lstrip('*').strip())
            p.style = 'List Bullet'
        
        # Listas numeradas
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text)
            p.style = 'List Number'
        
        # Párrafos normales
        elif line.strip():
            # Procesar formato inline (bold, italic, code)
            text = line.strip()
            p = doc.add_paragraph()
            
            # Procesar texto con formato
            last_end = 0
            
            # Buscar **texto** (bold)
            for match in re.finditer(r'\*\*(.+?)\*\*', text):
                p.add_run(text[last_end:match.start()])
                run = p.add_run(match.group(1))
                run.font.bold = True
                last_end = match.end()
            
            # Agregar resto del texto
            remaining = text[last_end:]
            
            # Buscar `código` (code)
            remaining_processed = ""
            last_end = 0
            for match in re.finditer(r'`([^`]+?)`', remaining):
                remaining_processed += remaining[last_end:match.start()]
                last_end = match.end()
                code_run = p.add_run(match.group(1))
                code_run.font.name = 'Courier New'
                code_run.font.color.rgb = RGBColor(192, 0, 0)
            
            remaining_processed += remaining[last_end:]
            if remaining_processed:
                p.add_run(remaining_processed)
        
        i += 1
    
    # Guardar documento
    doc.save(docx_file)
    print(f"✅ Documento creado exitosamente: {docx_file}")

if __name__ == '__main__':
    md_file = 'MANUAL_USUARIO_BASE.md'
    docx_file = 'MANUAL_USUARIO_BASE.docx'
    
    if os.path.exists(md_file):
        parse_markdown_to_docx(md_file, docx_file)
    else:
        print(f"❌ Error: No se encontró el archivo {md_file}")
